import docx

from alletra_onboard.application.asbuilt import AsBuiltData, generate_asbuilt

_SAMPLE = AsBuiltData(
    name="MPB10K-E24U21-LZ", model="Alletra MP B10000", serial_no="SGHD45FF0Y",
    controller_nodes="2", os_version="10.5.51", drive_cages="1", cache_gb="512",
    nvme_ssd_disks="24 x 15.36 TB", raw_capacity="368.6 TiB", raid="RAID 6 (6+2)",
    host_ports="4 x 32Gb FC (0:3:1, 0:3:2, 1:3:1, 1:3:2)",
    mgmt_ip="10.64.154.225", netmask="255.255.248.0", gateway="10.64.159.254",
    ntp="NTP1.bgl1.example.net", dns="10.203.96.10, 10.203.96.9",
    inventory="Node 0  Assembly ...\nNode 1  Assembly ...",
    checkhealth="Checking System ...\nSystem is healthy",
)


def _read(path):
    doc = docx.Document(str(path))
    table = {r.cells[0].text.strip(): r.cells[1].text.strip() for t in doc.tables for r in t.rows}
    text = "\n".join(p.text for p in doc.paragraphs)
    return doc, table, text


def test_generate_asbuilt_unbranded_builds_table_and_sections(tmp_path):
    # template=None -> plain document, but all the data still renders
    out = generate_asbuilt(_SAMPLE, tmp_path / "asbuilt.docx", template=None)
    _doc, table, text = _read(out)
    assert table["Serial No"] == "SGHD45FF0Y"
    assert table["InServ IP Address"] == "10.64.154.225"
    assert table["DNS Servers"] == "10.203.96.10, 10.203.96.9"
    assert table["Host Ports"].startswith("4 x 32Gb FC")
    assert "Node 0  Assembly" in text          # inventory
    assert "System is healthy" in text         # checkhealth
    assert "Environment Overview" in text and "Alletra configuration" in text


def test_generate_asbuilt_uses_template_clears_its_body_keeps_branding(tmp_path):
    # Build a stand-in "house-style" template: a header (branding) + instructional body to be cleared.
    tpl = docx.Document()
    tpl.sections[0].header.paragraphs[0].text = "HPE-BRAND-HEADER"
    tpl.add_paragraph("DELETE-ME instructional body from the template")
    tpl_path = tmp_path / "template.docx"
    tpl.save(str(tpl_path))

    out = generate_asbuilt(_SAMPLE, tmp_path / "branded.docx", template=tpl_path)
    doc, table, text = _read(out)
    assert "DELETE-ME" not in text                      # template's body content was cleared
    assert table["Serial No"] == "SGHD45FF0Y"           # our content rendered
    header = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "HPE-BRAND-HEADER" in header                 # branding (header) preserved


def test_generate_asbuilt_missing_values_render_as_dash(tmp_path):
    out = generate_asbuilt(AsBuiltData(serial_no="SGH123"), tmp_path / "sparse.docx", template=None)
    _doc, table, _text = _read(out)
    assert table["Serial No"] == "SGH123"
    assert table["Name"] == "-"   # empty -> dash, never a blank cell
