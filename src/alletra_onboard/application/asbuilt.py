"""As-built document generation (the 'Document' step — the LAST step after initialisation).

Fills HPE's "Block Storage" as-built Word template (bundled as ``resources/asbuilt_template.docx``)
with this deployment's **read-only** array facts, preserving everything the template already provides:
its cover page (with the ``<Customer Name>`` placeholder), table of contents, architecture sections,
HPE Graphik fonts, and Table 01's formatting. Only the per-deployment regions are written:

  * cover page — the customer name replaces ``<Customer Name>``.
  * Table 01 (Alletra configuration) — the config fields, matched by the label in column 0.
  * Alletra Inventory — ``showinventory`` output (verbatim, monospace — a raw hardware dump).
  * Alletra MP checkhealth output — rendered as two formatted Word tables (Summary + Details).

Parsing the raw ``show*`` text into ``AsBuiltData`` lives in ``asbuilt_parse.py`` (live-calibrated).
"""

from __future__ import annotations

import io
import os
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from alletra_onboard.application.asbuilt_docx import hpe_table

# Template Table-01 label (column 0, as it appears in the doc) -> AsBuiltData field. Match is
# whitespace-normalised + case-insensitive (the template mixes "InServ"/"Inserv").
_LABEL_TO_FIELD: dict[str, str] = {
    "name": "name",
    "model": "model",
    "serial no": "serial_no",
    "controller nodes": "controller_nodes",
    "os version": "os_version",
    "drive cages": "drive_cages",
    "cache(gb)": "cache_gb",
    "nvme ssd disks (15.36 tb)": "nvme_ssd_disks",
    "nvme ssd disks": "nvme_ssd_disks",
    "raw capacity": "raw_capacity",
    "raid": "raid",
    "host ports": "host_ports",
    "inserv ip address": "mgmt_ip",
    "inserv netmask address": "netmask",
    "inserv gateway address": "gateway",
    "ntp": "ntp",
    "dns servers": "dns",
}


@dataclass
class AsBuiltData:
    """The per-deployment values that fill the as-built (from read-only array reads + operator input)."""

    customer: str = ""       # operator-supplied — fills the cover page <Customer Name>
    site: str = ""           # operator-supplied — the array doesn't store a Location/Site
    name: str = ""
    model: str = ""
    serial_no: str = ""
    controller_nodes: str = ""
    os_version: str = ""
    drive_cages: str = ""
    cache_gb: str = ""
    nvme_ssd_disks: str = ""
    raw_capacity: str = ""
    raid: str = ""
    host_ports: str = ""
    mgmt_ip: str = ""
    netmask: str = ""
    gateway: str = ""
    ntp: str = ""
    dns: str = ""
    inventory: str = ""      # showinventory, verbatim
    checkhealth: str = ""    # checkhealth -svc -detail, verbatim


def _resource_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(getattr(sys, "_MEIPASS", ".")) / "alletra_onboard" / "resources"
    return Path(__file__).resolve().parent.parent / "resources"


def default_template() -> Path | None:
    """The as-built template: an explicit env override, else the bundled template (prefer the filled
    ``.docx`` over a bare ``.dotx``), else a ``templates/`` drop-in, else None (plain document)."""
    env = os.environ.get("ALLETRA_ASBUILT_TEMPLATE")
    if env and Path(env).is_file():
        return Path(env)
    rd = _resource_dir()
    for name in ("asbuilt_template.docx", "asbuilt_template.dotx"):
        if (rd / name).is_file():
            return rd / name
    bases = [Path.cwd()]
    if getattr(sys, "argv", None) and sys.argv[0]:
        bases.append(Path(sys.argv[0]).resolve().parent)
    for base in bases:
        for name in ("asbuilt_template.docx", "asbuilt_template.dotx"):
            if (base / "templates" / name).is_file():
                return base / "templates" / name
    return None


def _norm(text: str) -> str:
    return " ".join(str(text).split()).strip().rstrip(":").strip()


def _load_document(template: str | Path | None):
    """Open a .docx/.dotx as a python-docx Document (patching a .dotx content-type in memory), or a
    blank document when no template is given."""
    if template is None:
        return docx.Document()
    template = Path(template)
    if template.suffix.lower() == ".dotx":
        zin = zipfile.ZipFile(template)
        ct = zin.read("[Content_Types].xml").decode("utf-8").replace(
            "wordprocessingml.template.main+xml", "wordprocessingml.document.main+xml")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                zout.writestr(item, ct.encode("utf-8") if item == "[Content_Types].xml" else zin.read(item))
        zin.close()
        buf.seek(0)
        return docx.Document(buf)
    return docx.Document(str(template))


def _replace_text(doc, old: str, new: str) -> int:
    """Replace ``old`` with ``new`` across every text run in the body (incl. the cover-page sdt)."""
    count = 0
    for t in doc.element.body.iter(qn("w:t")):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            count += 1
    return count


def _set_cell_text(cell, text: str) -> None:
    first = cell.paragraphs[0]
    for run in list(first.runs):
        run._element.getparent().remove(run._element)
    first.add_run(text or "-")
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _insert_mono_after(heading, text: str):
    new_p = OxmlElement("w:p")
    heading._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, heading._parent)
    lines = (text or "").splitlines() or ["(no output captured)"]
    for i, line in enumerate(lines):
        run = para.add_run(line or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(7)
        if i < len(lines) - 1:
            run.add_break()
    return new_p


def _place_after(anchor, element):
    """Move a python-docx Paragraph/Table's XML to right after ``anchor`` (an lxml element)."""
    xml = element._p if hasattr(element, "_p") else element._tbl
    anchor.addnext(xml)
    return xml


def _content_widths(headers: list[str], rows) -> list[float]:
    """Column-width fractions proportional to the widest cell in each column (min 4% each)."""
    n = len(headers)
    widest = [max([len(headers[i])] + [len(r[i]) for r in rows if i < len(r)]) for i in range(n)]
    total = sum(widest) or n
    return [max(0.04, w / total) for w in widest]


def _bold_para_after(anchor, caption: str, doc):
    para = doc.add_paragraph()
    para.add_run(caption).bold = True
    return _place_after(anchor, para)


def _add_inventory_after(heading, text: str, doc) -> None:
    """Render each ``showinventory`` sub-section as its own HPE-styled table under the heading; falls
    back to a monospace dump if nothing parsed."""
    from alletra_onboard.application.asbuilt_parse import parse_inventory  # local: avoid import cycle

    sections = parse_inventory(text)
    if not sections:
        _insert_mono_after(heading, text)
        return
    anchor = heading._p
    for title, headers, rows in sections:
        if title:
            anchor = _bold_para_after(anchor, title, doc)
        tbl = hpe_table(doc, headers, rows, widths=_content_widths(headers, rows), font_size=7, header_size=7)
        anchor = _place_after(anchor, tbl)


def _add_checkhealth_after(heading, text: str, doc) -> None:
    """Render checkhealth as two HPE-styled Word tables (Summary + Details); falls back to monospace."""
    from alletra_onboard.application.asbuilt_parse import parse_checkhealth  # local: avoid import cycle

    summary, detail = parse_checkhealth(text)
    if not summary and not detail:
        _insert_mono_after(heading, text)
        return
    anchor = heading._p
    if summary:
        anchor = _bold_para_after(anchor, "Summary", doc)
        tbl = hpe_table(doc, ["Component", "Summary Description", "Qty"], summary,
                        widths=[0.20, 0.66, 0.14], font_size=8, header_size=8)
        anchor = _place_after(anchor, tbl)
    if detail:
        anchor = _bold_para_after(anchor, "Details", doc)
        tbl = hpe_table(doc, ["Component", "Identifier", "Detailed Description", "Resolution"], detail,
                        widths=[0.12, 0.28, 0.48, 0.12], font_size=8, header_size=8)
        anchor = _place_after(anchor, tbl)


def generate_asbuilt(data: AsBuiltData, out_path: str | Path, *, template: str | Path | None = ...) -> Path:
    """Fill the bundled template with ``data`` and write the finished as-built .docx to ``out_path``."""
    if template is ...:
        template = default_template()
    doc = _load_document(template)

    if data.customer:
        _replace_text(doc, "<Customer Name>", data.customer)

    values = {label: getattr(data, field) for label, field in _LABEL_TO_FIELD.items()}
    for table in doc.tables:
        for row in table.rows:
            key = _norm(row.cells[0].text).lower()
            if key in values and len(row.cells) > 1:
                _set_cell_text(row.cells[1], values[key])

    inv_heading = ch_heading = None
    for para in doc.paragraphs:
        if not para.style.name.lower().startswith("heading"):
            continue
        heading = _norm(para.text).lower()
        if heading == "alletra inventory":
            inv_heading = para
        elif "checkhealth" in heading:
            ch_heading = para
    if inv_heading is not None:
        _add_inventory_after(inv_heading, data.inventory, doc)
    if ch_heading is not None:
        _add_checkhealth_after(ch_heading, data.checkhealth, doc)

    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path
