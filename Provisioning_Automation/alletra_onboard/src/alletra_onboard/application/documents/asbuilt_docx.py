"""HPE-branded Word table helper for the as-built.

Ported from the operator's own ``sapdoc.py`` deliverable builder so the as-built's tables match the
house style of the other HPE deliverables: an HPE-green header row (white bold text), grey grid
borders, subtle zebra striping, tight cell margins, and a repeating header. ``hpe_table`` builds the
table at the end of the document and returns it, so the caller can reposition it (e.g. after a heading).
"""

from __future__ import annotations

from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

HEADER_FILL = "01A982"   # HPE green
ZEBRA_FILL = "F2F4F5"    # subtle zebra
GRID = "B1B9BE"          # grey border
WHITE = "FFFFFF"

_TBLPR_ORDER = ["w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual", "w:tblStyleRowBandSize",
                "w:tblStyleColBandSize", "w:tblW", "w:jc", "w:tblCellSpacing", "w:tblInd",
                "w:tblBorders", "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"]
_TCPR_ORDER = ["w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge", "w:tcBorders", "w:shd",
               "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"]


def _reorder(elem, order: list[str]) -> None:
    idx = {qn(t): i for i, t in enumerate(order)}
    for child in sorted(list(elem), key=lambda c: idx.get(c.tag, 999)):
        elem.append(child)


def _shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_margins(cell, top=40, bottom=40, left=100, right=100) -> None:
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    cell._tc.get_or_add_tcPr().append(m)


def _table_borders(tbl, color=GRID, sz=4) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tbl._tbl.tblPr.append(borders)


def _repeat_header(row) -> None:
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(th)


def _no_split(row) -> None:
    """Keep a row whole across a page boundary.

    The inventory and health tables regularly run past a page end; without this Word splits a row
    mid-height, leaving the top of the text on one page and the bottom on the next.
    """
    cs = OxmlElement("w:cantSplit")
    cs.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(cs)


def _content_width_emu(doc) -> int:
    sec = doc.sections[-1]
    return int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)


def _cell_text(cell, text, *, bold=False, color=None, size=8, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def hpe_table(doc, headers: list[str], rows, *, widths=None, font_size=8, header_size=8, zebra=True):
    """Build an HPE-styled table (green header + zebra) at the end of ``doc``; returns the table."""
    ncol = len(headers)
    tbl = doc.add_table(rows=1, cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    _table_borders(tbl)

    cw = _content_width_emu(doc)
    widths = widths or [1.0 / ncol] * ncol
    total = sum(widths) or 1
    col_emu = [int(cw * w / total) for w in widths]
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
            if i < len(col_emu):
                gc.set(qn("w:w"), str(int(col_emu[i] / 635)))
    _reorder(tbl._tbl.tblPr, _TBLPR_ORDER)

    header_row = tbl.rows[0]
    _repeat_header(header_row)
    _no_split(header_row)
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.width = Emu(col_emu[i])
        _shade(cell, HEADER_FILL)
        _cell_margins(cell)
        _cell_text(cell, text, bold=True, color=WHITE, size=header_size)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _reorder(cell._tc.get_or_add_tcPr(), _TCPR_ORDER)

    for ri, row in enumerate(rows):
        new_row = tbl.add_row()
        _no_split(new_row)
        cells = new_row.cells
        for i in range(ncol):
            cell = cells[i]
            cell.width = Emu(col_emu[i])
            _cell_margins(cell)
            if zebra and ri % 2 == 1:
                _shade(cell, ZEBRA_FILL)
            _cell_text(cell, row[i] if i < len(row) else "", size=font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _reorder(cell._tc.get_or_add_tcPr(), _TCPR_ORDER)
    return tbl
