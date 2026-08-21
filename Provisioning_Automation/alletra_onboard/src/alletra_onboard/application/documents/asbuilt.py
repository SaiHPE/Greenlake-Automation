"""As-built document generation (the 'Document' step — the LAST step after initialisation).

Fills HPE's "Block Storage" as-built Word template (bundled as ``resources/asbuilt_template.docx``)
with this deployment's **read-only** array facts, preserving everything the template already provides:
its cover page (with the ``<Customer Name>`` placeholder), table of contents, architecture sections,
HPE Graphik fonts, and Table 01's formatting. Only the per-deployment regions are written:

  * cover page — the customer name replaces ``<Customer Name>``.
  * Table 01 (Alletra configuration) — the config fields, matched by the label in column 0.
  * Alletra Inventory — ``showinventory`` output (verbatim, monospace — a raw hardware dump).
  * Alletra MP checkhealth output — rendered as two formatted Word tables (Summary + Details).

Parsing the raw ``show*`` text into ``AsBuiltData`` lives in ``asbuilt_parse.py`` (live-calibrated).
"""

from __future__ import annotations

import io
import os
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from alletra_onboard.application.documents.asbuilt_docx import hpe_table

# Template Table-01 label (column 0, as it appears in the doc) -> AsBuiltData field. Match is
# whitespace-normalised + case-insensitive (the template mixes "InServ"/"Inserv").
#: Narrative placeholders in the template -> the AsBuiltData field that fills them. Matched
#: case-insensitively; HPE's own template varies the capitalisation of the same field.
_PLACEHOLDER_TO_FIELD: dict[str, str] = {
    "<Customer Name>": "customer",
    "<Site Location>": "site",
    "<Application/Workload>": "application_workload",
    "<Purpose of using Block storage>": "purpose",
}
_PLACEHOLDER = re.compile(r"<[^<>]{1,60}>")

_LABEL_TO_FIELD: dict[str, str] = {
    "name": "name",
    "model": "model",
    "serial no": "serial_no",
    "controller nodes": "controller_nodes",
    "os version": "os_version",
    "drive cages": "drive_cages",
    "cache(gb)": "cache_gb",
    "nvme ssd disks (15.36 tb)": "nvme_ssd_disks",
    "nvme ssd disks": "nvme_ssd_disks",
    "raw capacity": "raw_capacity",
    "raid": "raid",
    "host ports": "host_ports",
    "inserv ip address": "mgmt_ip",
    "inserv netmask address": "netmask",
    "inserv gateway address": "gateway",
    "ntp": "ntp",
    "dns servers": "dns",
}


@dataclass
class AsBuiltData:
    """The per-deployment values that fill the as-built (from read-only array reads + operator input)."""

    customer: str = ""       # operator-supplied — fills <Customer Name> / <Customer name>
    site: str = ""           # operator-supplied — the array doesn't store a Location/Site
    application_workload: str = ""   # operator-supplied — <Application/Workload> (Introduction)
    purpose: str = ""                # operator-supplied — <Purpose of using Block storage>
    name: str = ""
    model: str = ""
    serial_no: str = ""
    controller_nodes: str = ""
    os_version: str = ""
    drive_cages: str = ""
    cache_gb: str = ""
    nvme_ssd_disks: str = ""
    raw_capacity: str = ""
    raid: str = ""
    host_ports: str = ""
    mgmt_ip: str = ""
    netmask: str = ""
    gateway: str = ""
    ntp: str = ""
    dns: str = ""
    inventory: str = ""      # showinventory, verbatim
    checkhealth: str = ""    # checkhealth -svc -detail, verbatim


def _resource_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(getattr(sys, "_MEIPASS", ".")) / "alletra_onboard" / "resources"
    return Path(__file__).resolve().parents[2] / "resources"  # …/src/alletra_onboard/resources


def default_template() -> Path | None:
    """The as-built template: an explicit env override, else the bundled template (prefer the filled
    ``.docx`` over a bare ``.dotx``), else a ``templates/`` drop-in, else None (plain document)."""
    env = os.environ.get("ALLETRA_ASBUILT_TEMPLATE")
    if env and Path(env).is_file():
        return Path(env)
    rd = _resource_dir()
    for name in ("asbuilt_template.docx", "asbuilt_template.dotx"):
        if (rd / name).is_file():
            return rd / name
    bases = [Path.cwd()]
    if getattr(sys, "argv", None) and sys.argv[0]:
        bases.append(Path(sys.argv[0]).resolve().parent)
    for base in bases:
        for name in ("asbuilt_template.docx", "asbuilt_template.dotx"):
            if (base / "templates" / name).is_file():
                return base / "templates" / name
    return None


def _norm(text: str) -> str:
    return " ".join(str(text).split()).strip().rstrip(":").strip()


def _load_document(template: str | Path | None):
    """Open a .docx/.dotx as a python-docx Document (patching a .dotx content-type in memory), or a
    blank document when no template is given."""
    if template is None:
        return docx.Document()
    template = Path(template)
    if template.suffix.lower() == ".dotx":
        zin = zipfile.ZipFile(template)
        ct = zin.read("[Content_Types].xml").decode("utf-8").replace(
            "wordprocessingml.template.main+xml", "wordprocessingml.document.main+xml")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                zout.writestr(item, ct.encode("utf-8") if item == "[Content_Types].xml" else zin.read(item))
        zin.close()
        buf.seek(0)
        return docx.Document(buf)
    return docx.Document(str(template))


def _text_parts(doc):
    """Every XML part whose text this document fills: the body AND each section's headers/footers.

    Word keeps running headers and footers in SEPARATE parts, so walking ``doc.element.body`` alone
    silently skips them — measured on HPE's template, the running header is
    ``"<Customer Name> HPE GreenLake for Block … Technical Whitepaper"`` and the footer carries
    ``<CustomerName>``, so every page after the cover shipped with a raw placeholder on it.

    Only parts this section actually OWNS are yielded: a header linked to the previous section has
    no definition of its own, and asking python-docx for one would create an empty part.
    """
    yield doc.element.body
    for section in doc.sections:
        for part in (
            section.first_page_header, section.header, section.even_page_header,
            section.first_page_footer, section.footer, section.even_page_footer,
        ):
            try:
                if part.is_linked_to_previous:
                    continue
                yield part._element
            except Exception:  # noqa: BLE001 - a malformed part must not stop the fill
                continue


def _placeholder_key(text: str) -> str:
    """``'<Customer Name>'``/``'<CustomerName>'``/``'<Customer name>'`` -> ``'customername'``.

    HPE's template spells one field FOUR ways across the cover, body, header and footer. Matching
    on the squeezed, lower-cased inner text makes the spelling irrelevant.
    """
    return re.sub(r"\s+", "", text.strip().strip("<>")).lower()


_FIELD_BY_KEY: dict[str, str] = {
    _placeholder_key(p): field for p, field in _PLACEHOLDER_TO_FIELD.items()
}


def _fill_placeholders(doc, data) -> int:
    """Replace every known ``<placeholder>`` with its value, across the body and headers/footers.

    Two passes, because Word splits text into runs at arbitrary points (revision marks,
    spell-check), so a placeholder can be *fragmented* across runs and never match a per-run
    search. Pass 1 replaces within each run — fast, and preserves any per-run formatting. Pass 2
    only touches paragraphs whose JOINED text still holds a known placeholder: it rewrites the
    paragraph's text into its first run, which is the price of repairing a split placeholder.
    """
    filled = 0

    def _value_for(match) -> str:
        field = _FIELD_BY_KEY.get(_placeholder_key(match.group(0)))
        if field:
            value = getattr(data, field, "")
            if value:
                return value
        return match.group(0)

    for part in _text_parts(doc):
        for node in part.iter(qn("w:t")):                      # pass 1: within a run
            if node.text and "<" in node.text:
                replaced = _PLACEHOLDER.sub(_value_for, node.text)
                if replaced != node.text:
                    node.text = replaced
                    filled += 1

        for para in part.iter(qn("w:p")):                      # pass 2: split across runs
            nodes = list(para.iter(qn("w:t")))
            if len(nodes) < 2:
                continue
            joined = "".join(n.text or "" for n in nodes)
            if "<" not in joined:
                continue
            if not any(_placeholder_key(m) in _FIELD_BY_KEY for m in _PLACEHOLDER.findall(joined)):
                continue
            replaced = _PLACEHOLDER.sub(_value_for, joined)
            if replaced == joined:
                continue
            nodes[0].text = replaced
            for extra in nodes[1:]:
                extra.text = ""
            filled += 1
    return filled


def _remaining_placeholders(doc) -> list[str]:
    """Every ``<...>`` still in the finished document — body, headers and footers — first seen first.

    Per-paragraph, not per-run: a placeholder fragmented across runs is still visible to the reader
    and must be reported, or the warning quietly under-counts what the customer will see.
    """
    found: list[str] = []
    for part in _text_parts(doc):
        for para in part.iter(qn("w:p")):
            joined = "".join(n.text or "" for n in para.iter(qn("w:t")))
            for match in _PLACEHOLDER.findall(joined):
                if match not in found:
                    found.append(match)
    return found


def _start_sections_on_new_pages(doc) -> int:
    """Begin every top-level section on a fresh page.

    The template carries exactly ONE explicit page break (before Executive Summary); every other
    Heading 1 flows inline, so once the generated inventory and health tables are injected the
    later headings land halfway down a page. ``page_break_before`` is the Word-native way to say
    this — it survives content reflow, unlike an inserted break paragraph that drifts as the tables
    above it grow. Headings that already carry an explicit break are left alone so no blank page
    appears, and each heading is kept with the content under it.
    """
    added = 0
    for para in doc.paragraphs:
        style = (para.style.name or "").lower() if para.style is not None else ""
        if not style.startswith("heading"):
            continue
        para.paragraph_format.keep_with_next = True          # never strand a heading at a page foot
        if style != "heading 1":
            continue
        if any(b.get(qn("w:type")) == "page" for b in para._p.iter(qn("w:br"))):
            continue                                          # the template already breaks here
        if para.paragraph_format.page_break_before:
            continue
        para.paragraph_format.page_break_before = True
        added += 1
    return added


def _drop_blank_paragraphs_after(heading) -> None:
    """Remove the run of empty paragraphs directly under a section heading.

    The template leaves ~23 blank paragraphs as space for the content we generate; without this the
    inserted tables are followed by a page of whitespace before the next heading.
    """
    node = heading._p.getnext()
    while node is not None and node.tag == qn("w:p"):
        following = node.getnext()
        if "".join(node.itertext()).strip():
            break
        node.getparent().remove(node)
        node = following


def _set_cell_text(cell, text: str) -> None:
    first = cell.paragraphs[0]
    for run in list(first.runs):
        run._element.getparent().remove(run._element)
    first.add_run(text or "-")
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _insert_mono_after(heading, text: str):
    new_p = OxmlElement("w:p")
    heading._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, heading._parent)
    lines = (text or "").splitlines() or ["(no output captured)"]
    for i, line in enumerate(lines):
        run = para.add_run(line or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(7)
        if i < len(lines) - 1:
            run.add_break()
    return new_p


def _place_after(anchor, element):
    """Move a python-docx Paragraph/Table's XML to right after ``anchor`` (an lxml element)."""
    xml = element._p if hasattr(element, "_p") else element._tbl
    anchor.addnext(xml)
    return xml


def _content_widths(headers: list[str], rows) -> list[float]:
    """Column-width fractions proportional to the widest cell in each column (min 4% each)."""
    n = len(headers)
    widest = [max([len(headers[i])] + [len(r[i]) for r in rows if i < len(r)]) for i in range(n)]
    total = sum(widest) or n
    return [max(0.04, w / total) for w in widest]


def _bold_para_after(anchor, caption: str, doc):
    para = doc.add_paragraph()
    para.add_run(caption).bold = True
    return _place_after(anchor, para)


def _add_inventory_after(heading, text: str, doc) -> None:
    """Render each ``showinventory`` sub-section as its own HPE-styled table under the heading; falls
    back to a monospace dump if nothing parsed."""
    from alletra_onboard.application.documents.asbuilt_parse import parse_inventory  # local: avoid import cycle

    sections = parse_inventory(text)
    if not sections:
        _insert_mono_after(heading, text)
        return
    anchor = heading._p
    for title, headers, rows in sections:
        if title:
            anchor = _bold_para_after(anchor, title, doc)
        tbl = hpe_table(doc, headers, rows, widths=_content_widths(headers, rows), font_size=7, header_size=7)
        anchor = _place_after(anchor, tbl)


def _add_checkhealth_after(heading, text: str, doc) -> None:
    """Render checkhealth as two HPE-styled Word tables (Summary + Details); falls back to monospace."""
    from alletra_onboard.application.documents.asbuilt_parse import parse_checkhealth  # local: avoid import cycle

    summary, detail = parse_checkhealth(text)
    if not summary and not detail:
        _insert_mono_after(heading, text)
        return
    anchor = heading._p
    if summary:
        anchor = _bold_para_after(anchor, "Summary", doc)
        tbl = hpe_table(doc, ["Component", "Summary Description", "Qty"], summary,
                        widths=[0.20, 0.66, 0.14], font_size=8, header_size=8)
        anchor = _place_after(anchor, tbl)
    if detail:
        anchor = _bold_para_after(anchor, "Details", doc)
        tbl = hpe_table(doc, ["Component", "Identifier", "Detailed Description", "Resolution"], detail,
                        widths=[0.12, 0.28, 0.48, 0.12], font_size=8, header_size=8)
        anchor = _place_after(anchor, tbl)


def generate_asbuilt(
    data: AsBuiltData, out_path: str | Path, *, template: str | Path | None = ...,
) -> tuple[Path, list[str]]:
    """Fill the bundled template with ``data`` and write the finished as-built .docx to ``out_path``.

    Returns ``(path, warnings)``. The document is ALWAYS produced — an operator in the field would
    rather hand over a partial document than be blocked — but anything that did not fill is
    reported so it can be seen before the document is sent:

      * a section whose heading could not be found (its content would be missing entirely)
      * any ``<placeholder>`` still present in the finished text

    Section headings are matched by INTENT, not by exact string. HPE renamed "Alletra Inventory" to
    "HPE GreenLake for Block hardware Inventory" between template revisions; the old exact match
    silently produced an as-built with no inventory section at all.
    """
    if template is ...:
        template = default_template()
    doc = _load_document(template)
    warnings: list[str] = []

    _fill_placeholders(doc, data)

    values = {label: getattr(data, field) for label, field in _LABEL_TO_FIELD.items()}
    for table in doc.tables:
        for row in table.rows:
            key = _norm(row.cells[0].text).lower()
            if key in values and len(row.cells) > 1:
                _set_cell_text(row.cells[1], values[key])

    inv_heading = ch_heading = None
    for para in doc.paragraphs:
        if not para.style.name.lower().startswith("heading"):
            continue
        heading = _norm(para.text).lower()
        if "inventory" in heading:
            inv_heading = para
        elif "checkhealth" in heading:
            ch_heading = para

    if inv_heading is not None:
        _drop_blank_paragraphs_after(inv_heading)
        _add_inventory_after(inv_heading, data.inventory, doc)
    else:
        warnings.append(
            "No 'Inventory' heading was found in the template, so the hardware inventory is NOT in "
            "this document."
        )
    if ch_heading is not None:
        _drop_blank_paragraphs_after(ch_heading)
        _add_checkhealth_after(ch_heading, data.checkhealth, doc)
    else:
        warnings.append(
            "No 'checkhealth' heading was found in the template, so the health report is NOT in "
            "this document."
        )

    _start_sections_on_new_pages(doc)

    left = _remaining_placeholders(doc)
    if left:
        warnings.append(
            "Placeholders are still unfilled and will be visible to the customer: "
            + ", ".join(left) + ". Fill them in the workbook (or the As-built step) and regenerate, "
            "or edit the document before sending."
        )

    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path, warnings
