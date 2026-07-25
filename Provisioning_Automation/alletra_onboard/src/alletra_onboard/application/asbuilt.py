"""As-built document generation (the 'Document' step — the LAST step after initialisation).

Fills HPE's "Block Storage" as-built Word template (bundled as ``resources/asbuilt_template.docx``)
with this deployment's **read-only** array facts, preserving everything the template already provides:
its cover page (with the ``<Customer Name>`` placeholder), table of contents, architecture sections,
HPE Graphik fonts, and Table 01's formatting. Only the per-deployment regions are written:

  * cover page — the customer name replaces ``<Customer Name>``.
  * Table 01 (Alletra configuration) — the config fields, matched by the label in column 0.
  * Alletra Inventory — ``showinventory`` output (verbatim, monospace — a raw hardware dump).
  * Alletra MP checkhealth output — rendered as two formatted Word tables (Summary + Details).

Parsing the raw ``show*`` text into ``AsBuiltData`` lives in ``asbuilt_parse.py`` (live-calibrated).
"""

from __future__ import annotations

import io
import os
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Template Table-01 label (column 0, as it appears in the doc) -> AsBuiltData field. Match is
# whitespace-normalised + case-insensitive (the template mixes "InServ"/"Inserv").
_LABEL_TO_FIELD: dict[str, str] = {
    "name": "name",
    "model": "model",
    "serial no": "serial_no",
    "controller nodes": "controller_nodes",
    "os version": "os_version",
    "drive cages": "drive_cages",
    "cache(gb)": "cache_gb",
    "nvme ssd disks (15.36 tb)": "nvme_ssd_disks",
    "nvme ssd disks": "nvme_ssd_disks",
    "raw capacity": "raw_capacity",
    "raid": "raid",
    "host ports": "host_ports",
    "inserv ip address": "mgmt_ip",
    "inserv netmask address": "netmask",
    "inserv gateway address": "gateway",
    "ntp": "ntp",
    "dns servers": "dns",
}


@dataclass
class AsBuiltData:
    """The per-deployment values that fill the as-built (from read-only array reads + operator input)."""

    customer: str = ""       # operator-supplied — fills the cover page <Customer Name>
    site: str = ""           # operator-supplied — the array doesn't store a Location/Site
    name: str = ""
    model: str = ""
    serial_no: str = ""
    controller_nodes: str = ""
    os_version: str = ""
    drive_cages: str = ""
    cache_gb: str = ""
    nvme_ssd_disks: str = ""
    raw_capacity: str = ""
    raid: str = ""
    host_ports: str = ""
    mgmt_ip: str = ""
    netmask: str = ""
    gateway: str = ""
    ntp: str = ""
    dns: str = ""
    inventory: str = ""      # showinventory, verbatim
    checkhealth: str = ""    # checkhealth -svc -detail, verbatim


def _resource_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(getattr(sys, "_MEIPASS", ".")) / "alletra_onboard" / "resources"
    return Path(__file__).resolve().parent.parent / "resources"


def default_template() -> Path | None:
    """The as-built template: an explicit env override, else the bundled template (prefer the filled
    ``.docx`` over a bare ``.dotx``), else a ``templates/`` drop-in, else None (plain document)."""
    env = os.environ.get("ALLETRA_ASBUILT_TEMPLATE")
    if env and Path(env).is_file():
        return Path(env)
    rd = _resource_dir()
    for name in ("asbuilt_template.docx", "asbuilt_template.dotx"):
        if (rd / name).is_file():
            return rd / name
    bases = [Path.cwd()]
    if getattr(sys, "argv", None) and sys.argv[0]:
        bases.append(Path(sys.argv[0]).resolve().parent)
    for base in bases:
        for name in ("asbuilt_template.docx", "asbuilt_template.dotx"):
            if (base / "templates" / name).is_file():
                return base / "templates" / name
    return None


def _norm(text: str) -> str:
    return " ".join(str(text).split()).strip().rstrip(":").strip()


def _load_document(template: str | Path | None):
    """Open a .docx/.dotx as a python-docx Document (patching a .dotx content-type in memory), or a
    blank document when no template is given."""
    if template is None:
        return docx.Document()
    template = Path(template)
    if template.suffix.lower() == ".dotx":
        zin = zipfile.ZipFile(template)
        ct = zin.read("[Content_Types].xml").decode("utf-8").replace(
            "wordprocessingml.template.main+xml", "wordprocessingml.document.main+xml")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                zout.writestr(item, ct.encode("utf-8") if item == "[Content_Types].xml" else zin.read(item))
        zin.close()
        buf.seek(0)
        return docx.Document(buf)
    return docx.Document(str(template))


def _replace_text(doc, old: str, new: str) -> int:
    """Replace ``old`` with ``new`` across every text run in the body (incl. the cover-page sdt)."""
    count = 0
    for t in doc.element.body.iter(qn("w:t")):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            count += 1
    return count


def _set_cell_text(cell, text: str) -> None:
    first = cell.paragraphs[0]
    for run in list(first.runs):
        run._element.getparent().remove(run._element)
    first.add_run(text or "-")
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _insert_mono_after(heading, text: str):
    new_p = OxmlElement("w:p")
    heading._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, heading._parent)
    lines = (text or "").splitlines() or ["(no output captured)"]
    for i, line in enumerate(lines):
        run = para.add_run(line or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(7)
        if i < len(lines) - 1:
            run.add_break()
    return new_p


def _table_style(doc) -> str | None:
    names = {s.name for s in doc.styles}
    for candidate in ("Grid Table 4 Accent 1", "Grid Table 4", "Table Grid", "Light Grid Accent 1", "Light List Accent 1"):
        if candidate in names:
            return candidate
    return None


def _place_after(anchor, element):
    """Move a python-docx Paragraph/Table's XML to right after ``anchor`` (an lxml element)."""
    xml = element._p if hasattr(element, "_p") else element._tbl
    anchor.addnext(xml)
    return xml


def _fill_row(table, values, *, bold: bool = False) -> None:
    cells = table.add_row().cells
    for i, value in enumerate(values):
        if i < len(cells):
            cell = cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.bold = bold
            run.font.size = Pt(8)


def _add_checkhealth_after(heading, text: str, doc) -> None:
    """Render checkhealth as two formatted Word tables (Summary + Details) after the heading; falls
    back to a monospace dump if the output couldn't be parsed into rows."""
    from alletra_onboard.application.asbuilt_parse import parse_checkhealth  # local: avoid import cycle

    summary, detail = parse_checkhealth(text)
    if not summary and not detail:
        _insert_mono_after(heading, text)
        return
    style = _table_style(doc)
    anchor = heading._p

    def label(caption: str) -> None:
        nonlocal anchor
        para = doc.add_paragraph()
        run = para.add_run(caption)
        run.bold = True
        anchor = _place_after(anchor, para)

    def table(headers: list[str], rows) -> None:
        nonlocal anchor
        tbl = doc.add_table(rows=0, cols=len(headers))
        if style:
            try:
                tbl.style = style
            except Exception:  # noqa: BLE001 - fall back to the default table style
                pass
        _fill_row(tbl, headers, bold=True)
        for row in rows:
            _fill_row(tbl, row)
        anchor = _place_after(anchor, tbl)

    if summary:
        label("Summary")
        table(["Component", "Summary Description", "Qty"], summary)
    if detail:
        label("Details")
        table(["Component", "Identifier", "Detailed Description", "Resolution"], detail)


def generate_asbuilt(data: AsBuiltData, out_path: str | Path, *, template: str | Path | None = ...) -> Path:
    """Fill the bundled template with ``data`` and write the finished as-built .docx to ``out_path``."""
    if template is ...:
        template = default_template()
    doc = _load_document(template)

    if data.customer:
        _replace_text(doc, "<Customer Name>", data.customer)

    values = {label: getattr(data, field) for label, field in _LABEL_TO_FIELD.items()}
    for table in doc.tables:
        for row in table.rows:
            key = _norm(row.cells[0].text).lower()
            if key in values and len(row.cells) > 1:
                _set_cell_text(row.cells[1], values[key])

    inv_heading = ch_heading = None
    for para in doc.paragraphs:
        if not para.style.name.lower().startswith("heading"):
            continue
        heading = _norm(para.text).lower()
        if heading == "alletra inventory":
            inv_heading = para
        elif "checkhealth" in heading:
            ch_heading = para
    if inv_heading is not None:
        _insert_mono_after(inv_heading, data.inventory)
    if ch_heading is not None:
        _add_checkhealth_after(ch_heading, data.checkhealth, doc)

    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path
