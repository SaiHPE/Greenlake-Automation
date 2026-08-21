import docx
from docx.oxml.ns import qn

from alletra_onboard.application.documents.asbuilt import AsBuiltData, default_template, generate_asbuilt

_CHECKHEALTH = (
    "Checking alert\n"
    "Component ----Summary Description---- Qty\n"
    "Alert     New alerts                    9\n"
    "Cage      Cages degraded                1\n"
    "----------------------------------------\n"
    "        2 total                        10\n"
    "\n"
    "Component --Identifier-- ---Detailed Description--- Resolution\n"
    "Alert     hw_cage:1      Cage over temperature      Manual\n"
    "-------------------------------------------------------------\n"
    "       1 total                                      10\n"
)

_INVENTORY = (  # showinventory -csvtable output
    "-----------------------Cage Inventory--------------------,,,\n"
    "Cage,-Name-,-Manufacturer-,-Type-\n"
    "1,cage1,HPE,DCN7\n"
    ",,,\n"
    "-----------------------Disk Inventory--------------------,,,\n"
    "Id,CagePos,State,--MFR--\n"
    "0,1:1,normal,SAMSUNG\n"
)

_SAMPLE = AsBuiltData(
    customer="ACME Corp", site="Bengaluru DC1",
    application_workload="VMware vSphere cluster",
    purpose="primary datastore for the vault zone",
    name="MPB10K-E24U21-LZ", model="HPE Alletra Storage MP",
    serial_no="SGHD45FF0Y", controller_nodes="2", os_version="10.5.51", drive_cages="1 (cage1)",
    cache_gb="503 GiB", nvme_ssd_disks="10 x 3.84 TB NVMe SSD", raw_capacity="34.9 TiB",
    raid="RAID 6 (SSD_r6)", host_ports="8 x 32Gbps FC target",
    mgmt_ip="10.64.154.225", netmask="255.255.248.0", gateway="10.64.159.254",
    ntp="ntp.example.net", dns="10.203.96.10, 10.203.96.9",
    inventory=_INVENTORY, checkhealth=_CHECKHEALTH,
)


def _read(path):
    doc = docx.Document(str(path))
    text = "".join(t.text or "" for t in doc.element.body.iter(qn("w:t")))
    return doc, text


def test_default_template_is_the_bundled_block_storage_docx():
    t = default_template()
    assert t is not None and t.name == "asbuilt_template.docx" and t.is_file()


def test_generate_fills_cover_table01_and_renders_checkhealth_tables(tmp_path):
    out, warnings = generate_asbuilt(_SAMPLE, tmp_path / "asbuilt.docx")
    doc, text = _read(out)

    # cover-page customer name replaced
    assert "ACME Corp" in text and "<Customer Name>" not in text
    assert warnings == [], warnings          # a fully-supplied deployment warns about nothing

    # Table 01 (the template's own table) filled by label
    t01 = {r.cells[0].text.strip(): r.cells[1].text.strip() for r in doc.tables[0].rows}
    assert t01["Serial No"] == "SGHD45FF0Y"
    assert t01["InServ IP Address"] == "10.64.154.225"
    assert t01["Cache(GB)"].startswith("503 GiB")

    # checkhealth + inventory rendered as HPE Word tables (found by their headers, any order)
    all_headers = [tuple(c.text.strip() for c in t.rows[0].cells) for t in doc.tables]
    assert ("Component", "Summary Description", "Qty") in all_headers
    assert any(h[:2] == ("Component", "Identifier") for h in all_headers)
    assert any("Manufacturer" in h for h in all_headers)  # an inventory sub-table

    # inventory values landed in a table, and the template's boilerplate is preserved
    assert "cage1" in text and "SAMSUNG" in text
    assert "Executive Summary" in text and "HPE InfoSight" in text


def test_generate_leaves_placeholder_when_no_customer(tmp_path):
    out, warnings = generate_asbuilt(AsBuiltData(serial_no="SGH123"), tmp_path / "no_cust.docx")
    _doc, text = _read(out)
    assert "<Customer Name>" in text  # untouched when the operator gives no customer
    assert any("Placeholders are still unfilled" in w for w in warnings)
    assert any("<Customer Name>" in w for w in warnings)


def test_every_narrative_placeholder_is_filled(tmp_path):
    """The 2026-08-17 template puts four operator-supplied values in prose, and spells the customer
    THREE ways — <Customer name> in Executive Summary, <Customer Name> in Introduction and
    Environment Overview. An exact-match replace shipped a literal placeholder in the opening
    paragraph of the document handed to the customer."""
    out, warnings = generate_asbuilt(_SAMPLE, tmp_path / "narrative.docx")
    _doc, text = _read(out)

    for value in ("ACME Corp", "Bengaluru DC1", "VMware vSphere cluster",
                  "primary datastore for the vault zone"):
        assert value in text, value
    assert "<" not in text.replace("<=", ""), "a placeholder survived into the finished document"
    assert warnings == []


def _all_text(doc):
    """Body + headers/footers — what the reader actually sees on every page."""
    from docx.oxml.ns import qn as _qn

    from alletra_onboard.application.documents.asbuilt import _text_parts

    return "".join(t.text or "" for part in _text_parts(doc) for t in part.iter(_qn("w:t")))


def test_customer_name_is_filled_in_the_running_header_and_footer(tmp_path):
    """MEASURED on the template: the running header is
    '<Customer Name> HPE GreenLake for Block … Technical Whitepaper' and the footer carries
    '<CustomerName>' — a FOURTH spelling, with no space. Word keeps both in separate XML parts, so
    walking doc.element.body left a raw placeholder on every page after the cover.
    """
    out, warnings = generate_asbuilt(_SAMPLE, tmp_path / "hdr.docx")
    doc = docx.Document(str(out))

    header = "".join(
        t.text or "" for s in doc.sections for t in s.header._element.iter(qn("w:t"))
    )
    footer = "".join(
        t.text or "" for s in doc.sections for t in s.footer._element.iter(qn("w:t"))
    )
    assert "ACME Corp" in header, f"header not filled: {header!r}"
    assert "<" not in header, f"placeholder left in header: {header!r}"
    assert "ACME Corp" in footer, f"footer not filled: {footer!r}"
    assert "<" not in footer, f"placeholder left in footer: {footer!r}"
    assert warnings == []


def test_spelling_variants_of_one_field_all_fill():
    """Four spellings of the same field across cover, body, header and footer."""
    from alletra_onboard.application.documents.asbuilt import _FIELD_BY_KEY, _placeholder_key

    for spelling in ("<Customer Name>", "<Customer name>", "<CustomerName>", "<customer  name>"):
        assert _FIELD_BY_KEY.get(_placeholder_key(spelling)) == "customer", spelling


def test_remaining_placeholders_sees_headers_and_split_runs(tmp_path):
    """An unfilled placeholder in a header must be WARNED about, not silently shipped."""
    out, warnings = generate_asbuilt(AsBuiltData(serial_no="SGH123"), tmp_path / "empty.docx")
    doc = docx.Document(str(out))
    header = "".join(t.text or "" for s in doc.sections for t in s.header._element.iter(qn("w:t")))

    assert "<" in header                                   # nothing to fill it with
    assert any("Placeholders are still unfilled" in w for w in warnings)


def test_every_top_level_section_starts_on_a_new_page(tmp_path):
    """The template has exactly ONE explicit page break; every other Heading 1 flowed inline, so
    the generated tables pushed later headings into the middle of a page."""
    out, _warnings = generate_asbuilt(_SAMPLE, tmp_path / "breaks.docx")
    doc = docx.Document(str(out))

    h1 = [p for p in doc.paragraphs if (p.style.name or "").lower() == "heading 1"]
    assert len(h1) >= 6
    for para in h1:
        explicit = any(b.get(qn("w:type")) == "page" for b in para._p.iter(qn("w:br")))
        assert para.paragraph_format.page_break_before or explicit, f"no break before {para.text!r}"
        assert para.paragraph_format.keep_with_next, f"heading can be stranded: {para.text!r}"


def test_generated_table_rows_do_not_split_across_pages(tmp_path):
    """A long inventory/health table must break BETWEEN rows, not through one."""
    out, _warnings = generate_asbuilt(_SAMPLE, tmp_path / "rows.docx")
    doc = docx.Document(str(out))

    generated = [t for t in doc.tables if "Component" in t.rows[0].cells[0].text or
                 "Cage" in t.rows[0].cells[0].text or "Id" in t.rows[0].cells[0].text]
    assert generated, "no generated tables found"
    for table in generated:
        for row in table.rows:
            trPr = row._tr.find(qn("w:trPr"))
            assert trPr is not None and trPr.find(qn("w:cantSplit")) is not None
        header_trPr = table.rows[0]._tr.find(qn("w:trPr"))
        assert header_trPr.find(qn("w:tblHeader")) is not None   # repeats on every page


def test_inventory_section_is_found_despite_the_heading_rename(tmp_path):
    """REGRESSION: the heading was matched with `== "alletra inventory"`. HPE renamed it to
    "HPE GreenLake for Block hardware Inventory", so the anchor missed and the ENTIRE hardware
    inventory was dropped from the document — with no error, no warning, and a saved .docx."""
    out, warnings = generate_asbuilt(_SAMPLE, tmp_path / "inv.docx")
    doc, text = _read(out)

    assert "HPE GreenLake for Block hardware Inventory" in text     # the new heading is present...
    assert "cage1" in text and "SAMSUNG" in text                     # ...and so is its content
    assert any("Manufacturer" in tuple(c.text.strip() for c in t.rows[0].cells) for t in doc.tables)
    assert not any("Inventory" in w for w in warnings)


def test_a_template_without_the_sections_warns_instead_of_silently_dropping_them(tmp_path):
    """Operator choice (2026-08-17): always produce the document, but never let a missing section
    pass unnoticed."""
    import docx as _docx

    bare = tmp_path / "bare_template.docx"
    stub = _docx.Document()
    stub.add_paragraph("Nothing useful here")
    stub.save(str(bare))

    out, warnings = generate_asbuilt(_SAMPLE, tmp_path / "bare.docx", template=bare)

    assert out.is_file()                                            # still generated
    assert any("Inventory" in w and "NOT in this document" in w for w in warnings)
    assert any("checkhealth" in w and "NOT in this document" in w for w in warnings)
